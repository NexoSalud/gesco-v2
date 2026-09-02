"""Router para Documentos de Contratista — subida, consulta y evaluación.

Rutas públicas (identificación por cédula del contratista):
  POST /api/v1/documentos/subir        — subir un documento (multipart)
  GET  /api/v1/documentos/contrato/{contrato_numero}?cedula=...  — listar docs de un contrato
  GET  /api/v1/documentos/{id}/descargar                          — descargar PDF

Rutas protegidas (requieren JWT):
  GET    /api/v1/documentos/admin/listar        — listar todos los docs con filtros
  PUT    /api/v1/documentos/{id}/evaluar        — aprobar/rechazar
"""

import io
import logging
import os
import uuid
import tempfile

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import Response
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.database import get_db
from app.models.contratista import Contratista
from app.models.contrato import Contrato
from app.models.documento_contratista import DocumentoContratista
from app.schemas.documento_contratista import (
    DocumentoContratistaCreate,
    DocumentoContratistaOut,
    DocumentoContratistaEvaluar,
    TIPOS_DOCUMENTO,
    TIPOS_DOCUMENTO_LABELS,
)
from app.routers.auth import get_current_user
from app.models.auth import Usuario
from app.models.periodo_evaluacion import PeriodoEvaluacion
from app.services.cuenta_cobro import generar_cuenta_cobro_docx

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/documentos", tags=["Documentos de Contratista"])

# Directorio para documentos contractuales
DOCS_DIR = "/app/uploads/documentos"
os.makedirs(DOCS_DIR, exist_ok=True)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".docx", ".doc"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
DOCX_EXTENSIONS = {".docx", ".doc"}


def _es_pdf_protegido(contenido: bytes) -> tuple[bool, str | None]:
    """Verifica si un PDF está protegido con contraseña usando pikepdf.

    Returns:
        (es_protegido, mensaje_error)
    """
    try:
        import pikepdf

        pdf = pikepdf.open(contenido)
        pdf.close()  # Se abrió bien → no tiene contraseña
        return False, None
    except pikepdf.PasswordError:
        return True, "El archivo PDF está protegido con contraseña. Quita la contraseña e intenta de nuevo."
    except Exception:
        # Si no se puede abrir con pikepdf, intentar validación básica
        pass
    return False, None


def _validar_pdf(contenido: bytes) -> None:
    """Valida que el contenido sea un PDF válido y sin contraseña."""
    # Verificar firma PDF
    if not contenido.startswith(b"%PDF"):
        raise HTTPException(400, "El archivo no es un PDF válido. Solo se aceptan archivos PDF.")

    # Verificar PDF protegido
    protegido, msg = _es_pdf_protegido(contenido)
    if protegido:
        raise HTTPException(400, msg)


def _meses_ejecucion(contrato) -> int:
    """Meses de ejecución del contrato, de fecha_inicio a fecha_fin (inclusive).

    Fallbacks: cuotas_total si no hay fechas; 6 meses por defecto.
    """
    if contrato.fecha_inicio and contrato.fecha_fin:
        meses = (
            (contrato.fecha_fin.year - contrato.fecha_inicio.year) * 12
            + (contrato.fecha_fin.month - contrato.fecha_inicio.month)
            + 1
        )
        if meses > 0:
            return meses
    if contrato.cuotas_total and contrato.cuotas_total > 0:
        return contrato.cuotas_total
    return 6


# ─── PÚBLICO: Rutas sin autenticación (solo validación por cédula) ──────────

@router.get("/cuenta-cobro/generar")
async def generar_cuenta_cobro(
    cedula: str = Query(..., min_length=1),
    contrato_numero: str = Query(...),
    periodo_id: int | None = Query(None),
    valor: float | None = Query(None, description="Valor a cobrar (opcional; si no viene, se calcula)"),
    db: AsyncSession = Depends(get_db),
):
    """Genera la cuenta de cobro en PDF con los datos contractuales del contratista
    y la guarda como documento (CUENTA_COBRO) del periodo, si aún no existe.
    Devuelve el PDF generado.
    """
    # Validar contratista por cédula
    result = await db.execute(
        select(Contratista).where(Contratista.identificacion == cedula)
    )
    contratista = result.scalar_one_or_none()
    if not contratista:
        raise HTTPException(404, "Contratista no encontrado con esa cédula")

    # Validar contrato
    result = await db.execute(
        select(Contrato).where(
            Contrato.numero_contrato == contrato_numero,
            Contrato.contratista_id == contratista.id,
        )
    )
    contrato = result.scalar_one_or_none()
    if not contrato:
        raise HTTPException(404, "Contrato no encontrado o no pertenece al contratista")

    # Resolver periodo
    periodo = None
    if periodo_id:
        p_res = await db.execute(
            select(PeriodoEvaluacion).where(PeriodoEvaluacion.id == periodo_id)
        )
        periodo = p_res.scalar_one_or_none()

    # Nombre del periodo (mes) para el título de la cuenta
    if periodo:
        periodo_nombre = periodo.nombre
    else:
        from datetime import date as _date
        from app.routers.evaluacion import MESES_ES as _MESES
        mes = contrato.fecha_inicio.month if contrato.fecha_inicio else _date.today().month
        anio = contrato.fecha_inicio.year if contrato.fecha_inicio else _date.today().year
        periodo_nombre = f"{_MESES[mes - 1]} {anio}"

    # Calcular honorarios: valor del contrato / meses de ejecución
    if valor is not None and valor > 0:
        valor_cobro = valor
    else:
        meses = _meses_ejecucion(contrato)
        valor_cobro = round(contrato.monto_total / meses, 2) if meses > 0 else contrato.monto_total

    # Número de la cuenta de cobro (conteo de CUENTA_COBRO existentes + 1)
    count_res = await db.execute(
        select(func.count(DocumentoContratista.id)).where(
            DocumentoContratista.contratista_id == contratista.id,
            DocumentoContratista.contrato_numero == contrato_numero,
            DocumentoContratista.tipo_documento == "CUENTA_COBRO",
        )
    )
    numero_cuenta_cobro = f"{(count_res.scalar() or 0) + 1:02d}"

    # Generar DOCX
    docx_bytes = generar_cuenta_cobro_docx(
        contratista_nombre=contratista.nombre,
        contratista_cedula=contratista.identificacion,
        expedida_en=contratista.expedida_en,
        banco=contratista.banco,
        tipo_cuenta=contratista.tipo_cuenta,
        numero_cuenta=contratista.numero_cuenta,
        numero_contrato=contrato.numero_contrato,
        objeto=contrato.objeto or "",
        valor=valor_cobro,
        periodo_nombre=periodo_nombre,
        numero=numero_cuenta_cobro,
    )

    # Guardar como documento si no existe uno ya para ese contrato+periodo
    existing = await db.execute(
        select(DocumentoContratista).where(
            DocumentoContratista.contratista_id == contratista.id,
            DocumentoContratista.contrato_numero == contrato_numero,
            DocumentoContratista.tipo_documento == "CUENTA_COBRO",
        )
    )
    if periodo_id:
        existing = await db.execute(
            select(DocumentoContratista).where(
                DocumentoContratista.contratista_id == contratista.id,
                DocumentoContratista.contrato_numero == contrato_numero,
                DocumentoContratista.tipo_documento == "CUENTA_COBRO",
                DocumentoContratista.periodo_id == periodo_id,
            )
        )

    doc_existente = existing.scalar_one_or_none()
    if doc_existente:
        # Re-generar archivo y mantener estado
        safe_name = os.path.basename(doc_existente.archivo_ruta)
        file_path = os.path.join(DOCS_DIR, safe_name)
        with open(file_path, "wb") as f:
            f.write(docx_bytes)
        doc_existente.estado = "PENDIENTE"
        await db.commit()
    else:
        safe_name = f"{uuid.uuid4()}.docx"
        file_path = os.path.join(DOCS_DIR, safe_name)
        with open(file_path, "wb") as f:
            f.write(docx_bytes)
        doc = DocumentoContratista(
            contratista_id=contratista.id,
            contrato_numero=contrato_numero,
            tipo_documento="CUENTA_COBRO",
            periodo_id=periodo_id,
            archivo_ruta=f"/uploads/documentos/{safe_name}",
            archivo_nombre=f"CUENTA_DE_COBRO_{numero_cuenta_cobro}_{periodo_nombre.replace(' ', '_')}.docx",
            archivo_tamano=len(docx_bytes),
            estado="PENDIENTE",
        )
        db.add(doc)
        await db.commit()
        logger.info(f"Cuenta de cobro generada: {contrato_numero} — contratista {contratista.id}, periodo {periodo_nombre}")

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=CUENTA_DE_COBRO_{numero_cuenta_cobro}_{periodo_nombre.replace(' ', '_')}.docx"
        },
    )


@router.post("/subir", response_model=DocumentoContratistaOut, status_code=201)
async def subir_documento(
    contratista_id: int = Form(...),
    contrato_numero: str = Form(...),
    tipo_documento: str = Form(...),
    periodo_id: int | None = Form(None),
    archivo: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Sube un documento contractual (PDF) para un contratista."""
    # Validar tipo de documento
    if tipo_documento not in TIPOS_DOCUMENTO:
        raise HTTPException(400, f"Tipo de documento inválido. Válidos: {', '.join(TIPOS_DOCUMENTO)}")

    # Validar extensión
    ext = ""
    if archivo.filename and "." in archivo.filename:
        ext = "." + archivo.filename.rsplit(".", 1)[-1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Formato no válido. Aceptados: PDF, JPG, PNG, DOCX")

    # Validar que el contratista existe
    result = await db.execute(
        select(Contratista).where(Contratista.id == contratista_id)
    )
    contratista = result.scalar_one_or_none()
    if not contratista:
        raise HTTPException(404, "Contratista no encontrado")

    # Validar que el contrato existe y pertenece al contratista
    result = await db.execute(
        select(Contrato).where(
            Contrato.numero_contrato == contrato_numero,
            Contrato.contratista_id == contratista_id,
        )
    )
    contrato = result.scalar_one_or_none()
    if not contrato:
        raise HTTPException(404, "Contrato no encontrado o no pertenece al contratista")

    # Leer contenido
    content = await archivo.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, f"El archivo excede el tamaño máximo de {MAX_FILE_SIZE // (1024*1024)}MB")

    # Convertir a PDF según tipo de archivo
    original_filename = archivo.filename or "documento"
    safe_name = f"{uuid.uuid4()}.pdf"
    file_path = os.path.join(DOCS_DIR, safe_name)

    if ext == ".pdf":
        # Validar PDF (formato + sin contraseña)
        _validar_pdf(content)
        with open(file_path, "wb") as f:
            f.write(content)

    elif ext in DOCX_EXTENSIONS:
        # Convertir DOCX a PDF usando python-docx + weasyprint
        try:
            from docx import Document as DocxDocument
            from weasyprint import HTML

            docx_doc = DocxDocument(io.BytesIO(content))
            html_parts = ["""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
                body { font-family: 'Times New Roman', serif; font-size: 12pt; margin: 2.5cm; }
                table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                td, th { border: 1px solid #333; padding: 4px 6px; font-size: 10pt; }
                th { background: #1a3a5c; color: white; }
            </style></head><body>
"""]

            for para in docx_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if para.style.name.startswith("Heading 1"):
                    html_parts.append(f"<h1>{text}</h1>")
                elif para.style.name.startswith("Heading 2"):
                    html_parts.append(f"<h2>{text}</h2>")
                elif para.style.name.startswith("Heading 3"):
                    html_parts.append(f"<h3>{text}</h3>")
                else:
                    html_parts.append(f"<p>{text}</p>")

            for table in docx_doc.tables:
                html_parts.append("<table>")
                for i, row in enumerate(table.rows):
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        tag = "th" if i == 0 else "td"
                        html_parts.append(f"<{tag}>{cell.text}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

            html_parts.append("</body></html>")
            html_str = "\n".join(html_parts)
            pdf_bytes = HTML(string=html_str).write_pdf()
            content = pdf_bytes
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logger.error(f"Error convirtiendo DOCX a PDF: {e}")
            raise HTTPException(500, f"Error al convertir el archivo DOCX a PDF: {str(e)}")

    elif ext in IMAGE_EXTENSIONS:
        # Convertir imagen a PDF
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(content))
            # Convertir a RGB si es RGBA o P
            if img.mode in ("RGBA", "P", "LA", "PA"):
                img = img.convert("RGB")
            pdf_buffer = io.BytesIO()
            img.save(pdf_buffer, format="PDF")
            content = pdf_buffer.getvalue()
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logger.error(f"Error convirtiendo imagen a PDF: {e}")
            raise HTTPException(500, f"Error al convertir la imagen a PDF: {str(e)}")

    # Crear registro en BD
    doc = DocumentoContratista(
        contratista_id=contratista_id,
        contrato_numero=contrato_numero,
        tipo_documento=tipo_documento,
        periodo_id=periodo_id,
        archivo_ruta=f"/uploads/documentos/{safe_name}",
        archivo_nombre=original_filename,
        archivo_tamano=len(content),
        estado="PENDIENTE",
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    logger.info(f"Documento subido: {tipo_documento} — contratista {contratista_id}, contrato {contrato_numero}")

    return DocumentoContratistaOut(
        id=doc.id,
        contratista_id=doc.contratista_id,
        contrato_numero=doc.contrato_numero,
        tipo_documento=doc.tipo_documento,
        archivo_ruta=doc.archivo_ruta,
        archivo_nombre=doc.archivo_nombre,
        archivo_tamano=doc.archivo_tamano,
        estado=doc.estado,
        observacion=doc.observacion,
        created_at=doc.created_at,
        updated_at=doc.updated_at,
        evaluated_at=doc.evaluated_at,
        contratista_nombre=contratista.nombre,
        contratista_identificacion=contratista.identificacion,
    )


@router.get("/contrato/{contrato_numero}", response_model=list[DocumentoContratistaOut])
async def listar_documentos_contrato(
    contrato_numero: str,
    cedula: str = Query(..., min_length=1),
    periodo_id: int | None = Query(None),
    db: AsyncSession = Depends(get_db),
):
    """Lista los documentos de un contrato. Requiere cédula del contratista como validación."""
    # Validar contratista por cédula
    result = await db.execute(
        select(Contratista).where(Contratista.identificacion == cedula)
    )
    contratista = result.scalar_one_or_none()
    if not contratista:
        raise HTTPException(404, "Contratista no encontrado con esa cédula")

    # Buscar documentos del contratista para ese contrato
    stmt = (
        select(DocumentoContratista)
        .where(
            DocumentoContratista.contratista_id == contratista.id,
            DocumentoContratista.contrato_numero == contrato_numero,
        )
        .order_by(DocumentoContratista.tipo_documento, DocumentoContratista.created_at.desc())
    )
    if periodo_id is not None:
        stmt = stmt.where(DocumentoContratista.periodo_id == periodo_id)
    result = await db.execute(stmt)
    docs = result.scalars().all()

    return [
        DocumentoContratistaOut(
            id=doc.id,
            contratista_id=doc.contratista_id,
            contrato_numero=doc.contrato_numero,
            tipo_documento=doc.tipo_documento,
            archivo_ruta=doc.archivo_ruta,
            archivo_nombre=doc.archivo_nombre,
            archivo_tamano=doc.archivo_tamano,
            estado=doc.estado,
            observacion=doc.observacion,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
            evaluated_at=doc.evaluated_at,
            contratista_nombre=contratista.nombre,
            contratista_identificacion=contratista.identificacion,
        )
        for doc in docs
    ]


# ─── PROTEGIDO: Dashboard (coordinadora / administrativo) ────────────────────

@router.get("/admin/listar", response_model=list[DocumentoContratistaOut])
async def listar_todos_documentos(
    contratista_id: int | None = Query(None),
    contrato_numero: str | None = Query(None),
    tipo_documento: str | None = Query(None),
    estado: str | None = Query(None),
    periodo_id: int | None = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: Usuario = Depends(get_current_user),
):
    """Lista todos los documentos con filtros opcionales. Requiere autenticación."""
    stmt = select(DocumentoContratista).order_by(DocumentoContratista.created_at.desc())

    if contratista_id:
        stmt = stmt.where(DocumentoContratista.contratista_id == contratista_id)
    if contrato_numero:
        stmt = stmt.where(DocumentoContratista.contrato_numero == contrato_numero)
    if tipo_documento:
        stmt = stmt.where(DocumentoContratista.tipo_documento == tipo_documento)
    if estado:
        stmt = stmt.where(DocumentoContratista.estado == estado)
    if periodo_id is not None:
        stmt = stmt.where(DocumentoContratista.periodo_id == periodo_id)

    result = await db.execute(stmt)
    docs = result.scalars().all()

    out = []
    for doc in docs:
        # Obtener nombre del contratista
        cont_result = await db.execute(
            select(Contratista.nombre, Contratista.identificacion).where(
                Contratista.id == doc.contratista_id
            )
        )
        cont_row = cont_result.one_or_none()
        cont_nombre = cont_row[0] if cont_row else None
        cont_ident = cont_row[1] if cont_row else None

        out.append(DocumentoContratistaOut(
            id=doc.id,
            contratista_id=doc.contratista_id,
            contrato_numero=doc.contrato_numero,
            tipo_documento=doc.tipo_documento,
            archivo_ruta=doc.archivo_ruta,
            archivo_nombre=doc.archivo_nombre,
            archivo_tamano=doc.archivo_tamano,
            estado=doc.estado,
            observacion=doc.observacion,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
            evaluated_at=doc.evaluated_at,
            contratista_nombre=cont_nombre,
            contratista_identificacion=cont_ident,
        ))

    return out


@router.put("/{documento_id}/evaluar", response_model=DocumentoContratistaOut)
async def evaluar_documento(
    documento_id: int,
    data: DocumentoContratistaEvaluar,
    db: AsyncSession = Depends(get_db),
    current_user: Usuario = Depends(get_current_user),
):
    """Aprobar o rechazar un documento. Requiere autenticación."""
    result = await db.execute(
        select(DocumentoContratista).where(DocumentoContratista.id == documento_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")

    try:
        doc.estado = data.estado
        doc.observacion = data.observacion
        doc.evaluated_by = current_user.id
        from datetime import datetime
        doc.evaluated_at = datetime.utcnow()

        await db.commit()
        await db.refresh(doc)
    except Exception as e:
        import traceback
        logger.error(f"Error en evaluar_documento (commit/refresh): {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, f"Error al evaluar documento: {e}")

    try:
        # Obtener datos del contratista
        cont_result = await db.execute(
            select(Contratista.nombre, Contratista.identificacion).where(
                Contratista.id == doc.contratista_id
            )
        )
        cont_row = cont_result.one_or_none()
        cont_nombre = cont_row[0] if cont_row else None
        cont_ident = cont_row[1] if cont_row else None

        logger.info(f"Documento {doc.id} evaluado como {data.estado} por usuario {current_user.id}")

        return DocumentoContratistaOut(
            id=doc.id,
            contratista_id=doc.contratista_id,
            contrato_numero=doc.contrato_numero,
            tipo_documento=doc.tipo_documento,
            archivo_ruta=doc.archivo_ruta,
            archivo_nombre=doc.archivo_nombre,
            archivo_tamano=doc.archivo_tamano,
            estado=doc.estado,
            observacion=doc.observacion,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
            evaluated_at=doc.evaluated_at,
            contratista_nombre=cont_nombre,
            contratista_identificacion=cont_ident,
        )
    except Exception as e:
        import traceback
        logger.error(f"Error en evaluar_documento (serialización): {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, f"Error al serializar documento: {e}")


@router.delete("/{documento_id}", status_code=204)
async def eliminar_documento(
    documento_id: int,
    db: AsyncSession = Depends(get_db),
):
    """Elimina un documento y su archivo del disco."""
    result = await db.execute(
        select(DocumentoContratista).where(DocumentoContratista.id == documento_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Documento no encontrado")

    # Eliminar archivo del disco
    if doc.archivo_ruta:
        file_path = os.path.join("/app/uploads", doc.archivo_ruta.replace("/uploads/", "", 1))
        # También probar ruta directa
        if not os.path.exists(file_path):
            file_path = doc.archivo_ruta.lstrip("/")
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Archivo eliminado: {file_path}")
        except Exception as e:
            logger.warning(f"No se pudo eliminar el archivo {file_path}: {e}")

    await db.delete(doc)
    await db.commit()
    return None


# ─── Utilidad: tipos de documento ────────────────────────────────────────────

@router.get("/tipos")
async def listar_tipos_documento():
    """Devuelve los tipos de documento disponibles con sus etiquetas."""
    return {
        "tipos": [
            {"valor": t, "etiqueta": TIPOS_DOCUMENTO_LABELS.get(t, t)}
            for t in TIPOS_DOCUMENTO
        ]
    }
