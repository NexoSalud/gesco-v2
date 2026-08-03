"""Servicio de generación de Informe de Evaluación de Cumplimiento.
Genera PDF (WeasyPrint) y DOCX (python-docx) con estilo profesional tipo documento formal.
"""

import io
import logging
import base64
import os
from datetime import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

logger = logging.getLogger(__name__)

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

MESES = [
    "", "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
    "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE",
]

MESES_TITULO = [
    "", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]

FONT_NAME = "Arial"
FONT_SIZE = Pt(10)
FONT_SIZE_TITLE = Pt(14)

COLOR_PRIMARY = RGBColor(0, 51, 102)
COLOR_TEXT = RGBColor(51, 51, 51)

# ─── Renderizado común ───────────────────────────────────────────────────

def _cargar_logo_base64() -> str:
    """Carga el logo de la ESE desde static/ y lo devuelve en base64."""
    logo_path = Path(__file__).parent.parent / "static" / "logo_es.png"
    if logo_path.exists():
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return ""


def _build_context(contratista: dict, contratos: list, resumen: dict) -> dict:
    """Construye el contexto unificado para PDF y DOCX."""
    today = datetime.now()
    fecha_informe = f"{today.day:02d} de {MESES[today.month]} de {today.year}"

    total_actividades = 0
    for c in contratos:
        for act in c.get("actividades", []):
            total_actividades += 1
            evs = act.get("evidencias", [])
            act["total_evidencias"] = len(evs)
            act["observacion"] = None
            if evs:
                act["estado_global"] = "APROBADO"
                # Ocultar observaciones cuando hay evidencias aprobadas
                for ev in evs:
                    ev["observacion_coordinadora"] = None
            else:
                act["estado_global"] = "SIN_EVIDENCIA"

    total_ev = resumen.get("total_actividades", total_actividades) or 1

    # Agrupar documentos contractuales por estado
    total_docs = 0
    docs_aprobados = 0
    docs_rechazados = 0
    docs_pendientes = 0
    for c in contratos:
        for doc in c.get("documentos", []):
            total_docs += 1
            if doc["estado"] == "APROBADO":
                docs_aprobados += 1
            elif doc["estado"] == "RECHAZADO":
                docs_rechazados += 1
            else:
                docs_pendientes += 1

    return {
        "logo_base64": _cargar_logo_base64(),
        "nombre": contratista.get("nombre", ""),
        "identificacion": contratista.get("identificacion", ""),
        "telefono": contratista.get("telefono"),
        "correo": contratista.get("correo"),
        "contrato": contratos[0]["numero_contrato"] if contratos else "",
        "perfil": contratos[0].get("perfil") if contratos else "",
        "periodo": MESES[today.month],
        "fecha_informe": fecha_informe,
        "total_actividades": total_actividades,
        "aprobadas": resumen.get("aprobadas", 0),
        "rechazadas": resumen.get("rechazadas", 0),
        "pendientes": resumen.get("pendientes", 0),
        "sin_evidencia": resumen.get("sin_evidencia", 0),
        "porcentaje": resumen.get("porcentaje_cumplimiento", 0),
        "pct_aprobadas": round(resumen.get("aprobadas", 0) / total_ev * 100, 1),
        "pct_rechazadas": round(resumen.get("rechazadas", 0) / total_ev * 100, 1),
        "pct_pendientes": round(resumen.get("pendientes", 0) / total_ev * 100, 1),
        "pct_sin_evidencia": round(resumen.get("sin_evidencia", 0) / total_ev * 100, 1),
        "contratos": contratos,
        "observaciones": [],
        "documentos": contratos[0].get("documentos", []) if contratos else [],
        "total_documentos": total_docs,
        "documentos_aprobados": docs_aprobados,
        "documentos_rechazados": docs_rechazados,
        "documentos_pendientes": docs_pendientes,
    }


# ─── Generador PDF ────────────────────────────────────────────────────────

def _recolectar_anexos(contratos: list) -> list[dict]:
    """Recolecta los PDFs de documentos contractuales aprobados para anexar."""
    anexos = []
    logger.info(f"_recolectar_anexos: {len(contratos)} contratos")
    for idx, c in enumerate(contratos):
        docs = c.get("documentos", [])
        logger.info(f"  Contrato {idx}: {len(docs)} documentos")
        for doc in docs:
            estado = doc.get("estado", "?")
            tipo = doc.get("tipo_documento", "?")
            archivo_ruta = doc.get("archivo_ruta", "")
            logger.info(f"    Doc: {tipo} estado={estado} ruta={archivo_ruta!r}")
            if estado != "APROBADO":
                continue
            if not archivo_ruta:
                continue
            pdf_path = archivo_ruta.lstrip("/")
            if pdf_path.startswith("uploads/"):
                pdf_path = os.path.join("/app", pdf_path)
            else:
                pdf_path = os.path.join("/app/uploads", pdf_path)
            exists = os.path.exists(pdf_path)
            logger.info(f"    -> resolved={pdf_path!r} exists={exists}")
            if not exists:
                # Try alternative: maybe the file is directly in /app/uploads/
                alt_path = os.path.join("/app", archivo_ruta.lstrip("/"))
                logger.info(f"    -> alt={alt_path!r} exists={os.path.exists(alt_path)}")
                if os.path.exists(alt_path):
                    pdf_path = alt_path
                    exists = True
            if exists:
                anexos.append({
                    "ruta": pdf_path,
                    "tipo": tipo,
                    "nombre": doc.get("archivo_nombre", "documento"),
                })
                logger.info(f"    -> AGREGADO como anexo")
    logger.info(f"_recolectar_anexos: {len(anexos)} anexos encontrados")
    return anexos


def generar_pdf(contratista: dict, contratos: list, resumen: dict) -> bytes:
    """Genera el PDF del informe de evaluación con documentos contractuales anexados."""
    ctx = _build_context(contratista, contratos, resumen)

    env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
    template = env.get_template("informe_evaluacion.html")
    html_str = template.render(**ctx)

    # Generar PDF principal
    pdf_bytes = HTML(string=html_str).write_pdf()

    # Anexar documentos contractuales aprobados
    anexos = _recolectar_anexos(contratos)
    logger.info(f"generar_pdf: {len(anexos)} anexos para merge")
    if not anexos:
        logger.info("generar_pdf: sin anexos, PDF sin documentos contractuales")
        return pdf_bytes

    try:
        from pikepdf import Pdf
        logger.info("generar_pdf: pikepdf importado correctamente")

        with Pdf.open(io.BytesIO(pdf_bytes)) as main_pdf:
            logger.info(f"generar_pdf: PDF principal tiene {len(main_pdf.pages)} páginas")
            for anexo in anexos:
                try:
                    with Pdf.open(anexo["ruta"]) as doc_pdf:
                        doc_pages = len(doc_pdf.pages)
                        main_pdf.pages.extend(doc_pdf.pages)
                        logger.info(f"Anexado: {anexo['nombre']} ({doc_pages} páginas)")
                except Exception as e:
                    logger.warning(f"No se pudo anexar {anexo['nombre']}: {e}")
                    continue

            output = io.BytesIO()
            main_pdf.save(output)
            result = output.getvalue()
            logger.info(f"PDF generado con {len(anexos)} anexos contractuales ({len(result)} bytes)")
            return result
    except ImportError as e:
        logger.warning(f"pikepdf no disponible ({e}) — se omite anexado de documentos")
        return pdf_bytes
    except Exception as e:
        logger.error(f"Error anexando documentos contractuales: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return pdf_bytes


# ─── Generador DOCX ───────────────────────────────────────────────────────

def _coerce_size(size):
    """Convierte size a Length sin doble conversión.

    Si ya es un Length (p.ej. Pt(10)), se usa tal cual; si es un número,
    se interpreta como puntos. Evita el bug de Pt(Pt(10)) que produce
    tamaños de fuente gigantes (una letra por página).
    """
    if size is None:
        return None
    return size if isinstance(size, Length) else Pt(size)


def _add_styled_paragraph(doc, text, style_name=None, bold=False, size=None, color=None, alignment=None, space_after=None):
    """Agrega un párrafo con estilo."""
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = _coerce_size(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_cell_text(cell, text, bold=False, size=9, color=None, alignment=None):
    """Agrega texto formateado a una celda de tabla."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = _coerce_size(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return cell


def _set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


# ─── Helpers para el nuevo formato INFORME DE ACTIVIDADES ──────────────

def _render_html_to_cell(cell, html_text: str, font_size=9):
    """Renderiza HTML en una celda con formato (bold, saltos de línea, tablas).

    Si el contenido NO es HTML (sin tags), lo escribe como texto plano.
    Si tiene tags HTML, los renderiza con formato real en la celda.
    """
    if not html_text or ('<' not in html_text):
        # Texto plano — escribir directo
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(str(html_text))
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"
        return

    import re as _re2

    # Limpiar la celda
    for p in cell.paragraphs:
        p.clear()

    # Parsear HTML en tokens: texto plano, <br>, <b>/<strong>, <i>/<em>,
    # <p>, </p>, <li>, </li>, <table>...</table>
    # Simplificación: convertir a texto con marcas, luego construir paragraphs

    text = html_text
    # 1. Reemplazar <br> → salto de línea
    text = _re2.sub(r'<br\s*/?>', '\n', text, flags=_re2.IGNORECASE)
    # 2. </p> → doble salto
    text = _re2.sub(r'</p>', '\n\n', text, flags=_re2.IGNORECASE)
    # 3. <li> → bullet
    text = _re2.sub(r'<li[^>]*>', '• ', text, flags=_re2.IGNORECASE)
    text = _re2.sub(r'</li>', '\n', text, flags=_re2.IGNORECASE)
    # 4. <td> y <th> → tab
    text = _re2.sub(r'</t[dh]>', '\t', text, flags=_re2.IGNORECASE)
    text = _re2.sub(r'</tr>', '\n', text, flags=_re2.IGNORECASE)

    # 5. Extraer contenido de tablas para renderizarlas como tablas DOCX reales
    table_blocks = list(_re2.finditer(
        r'<table[^>]*>(.*?)</table>', text, _re2.DOTALL | _re2.IGNORECASE))
    # Reemplazar tablas por placeholder para procesar después
    table_placeholders = []
    for match in table_blocks:
        table_placeholders.append(match.group(0))
    text = _re2.sub(
        r'<table[^>]*>.*?</table>', '<<TABLE_PLACEHOLDER>>',
        text, flags=_re2.DOTALL | _re2.IGNORECASE)

    # 6. Ahora procesar el texto con formato inline (bold, italic)
    # Simplificación: tokenizar en (is_bold, text) para cada línea
    lines = text.split('\n')
    ph_idx = 0

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Si es un placeholder de tabla
        if '<<TABLE_PLACEHOLDER>>' in line and ph_idx < len(table_placeholders):
            table_html = table_placeholders[ph_idx]
            ph_idx += 1
            # Renderizar tabla simple
            rows_html = _re2.findall(
                r'<tr[^>]*>(.*?)</tr>', table_html,
                _re2.DOTALL | _re2.IGNORECASE)
            if rows_html:
                num_cols = max(
                    len(_re2.findall(r'<t[dh][^>]*>.*?</t[dh]>',
                                     rh, _re2.DOTALL | _re2.IGNORECASE))
                    for rh in rows_html)
                tbl = cell.add_table(rows=0, cols=max(num_cols, 1))
                _set_table_borders(tbl, sz="4", color="999999")
                for rh in rows_html:
                    cells_text = _re2.findall(
                        r'<t[dh][^>]*>(.*?)</t[dh]>',
                        rh, _re2.DOTALL | _re2.IGNORECASE)
                    clean_cells = [_re2.sub(r'<[^>]+>', '', c).strip()
                                   for c in cells_text]
                    is_header = '<th' in rh.lower()
                    row = tbl.add_row()
                    for ci, ct in enumerate(clean_cells):
                        if ci < len(row.cells):
                            cell2 = row.cells[ci]
                            cell2.text = ""
                            p2 = cell2.paragraphs[0]
                            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            r2 = p2.add_run(ct)
                            r2.font.size = Pt(font_size - 1)
                            r2.font.name = "Times New Roman"
                            r2.bold = is_header
                            if is_header:
                                _set_cell_shading(cell2, "E8E8E8")
                # Párrafo vacío como separador después de tabla
                sp = cell.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
            continue

        # 7. Limpiar tags HTML restantes pero preservar bold/italic/strikethrough
        # Estrategia: procesar runs inline
        remaining = line

        # Añadir un paragraph por línea
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        # Parsear inline formatting tags
        pos = 0
        while pos < len(remaining):
            # Buscar próximo tag
            next_tag = _re2.search(
                r'<(/?)(b|strong|i|em|u|s|del|strike)\s*>',
                remaining[pos:], _re2.IGNORECASE)

            if not next_tag:
                # No más tags — añadir resto del texto
                plain = remaining[pos:]
                plain = _re2.sub(r'<[^>]+>', '', plain)  # limpiar otros tags
                plain = plain.replace('&nbsp;', ' ').replace('&amp;', '&')
                plain = plain.replace('&lt;', '<').replace('&gt;', '>')
                if plain.strip():
                    run = p.add_run(plain)
                    run.font.size = Pt(font_size)
                    run.font.name = "Times New Roman"
                break

            # Texto antes del tag
            before = remaining[pos:pos + next_tag.start()]
            before = _re2.sub(r'<[^>]+>', '', before)
            before = before.replace('&nbsp;', ' ').replace('&amp;', '&')
            before = before.replace('&lt;', '<').replace('&gt;', '>')
            if before.strip():
                run = p.add_run(before)
                run.font.size = Pt(font_size)
                run.font.name = "Times New Roman"

            # Avanzar posición
            tag_name = next_tag.group(2).lower()
            is_closing = next_tag.group(1) == '/'
            pos += next_tag.end()

            if is_closing:
                continue  # ya procesamos el formato en la apertura

            # Buscar texto hasta el cierre del tag
            close_pattern = _re2.compile(
                rf'</{tag_name}\s*>', _re2.IGNORECASE)
            close_match = close_pattern.search(remaining[pos:])
            if close_match:
                inner = remaining[pos:pos + close_match.start()]
                inner = _re2.sub(r'<[^>]+>', '', inner)
                inner = inner.replace('&nbsp;', ' ').replace('&amp;', '&')
                inner = inner.replace('&lt;', '<').replace('&gt;', '>')
                if inner.strip():
                    run = p.add_run(inner)
                    run.font.size = Pt(font_size)
                    run.font.name = "Times New Roman"
                    if tag_name in ('b', 'strong'):
                        run.bold = True
                    elif tag_name in ('i', 'em'):
                        run.italic = True
                    elif tag_name == 'u':
                        run.underline = True
                    elif tag_name in ('s', 'del', 'strike'):
                        run.font.strike = True
                pos = pos + close_match.end()
            else:
                # No se encontró cierre — tratar como texto normal
                pass

    # Si no se añadió nada (celda vacía), poner texto plano
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip():
        cell.paragraphs[0].clear()
        clean_text = _re2.sub(r'<[^>]+>', '', html_text)
        clean_text = clean_text.replace('&nbsp;', ' ').replace('&amp;', '&')
        clean_text = clean_text.replace('&lt;', '<').replace('&gt;', '>')
        clean_text = clean_text.strip()
        run = cell.paragraphs[0].add_run(clean_text)
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"

def _set_cell_margins(cell, top=28, bottom=28, left=57, right=57):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_table_borders(table, sz="4", color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_paragraph(doc, text, bold=False, size=10, color=COLOR_TEXT,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = _coerce_size(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = color
    return p


def _make_header_row(table, text, num_cols, bg="003366"):
    row = table.rows[0] if table.rows else table.add_row()
    if num_cols > 1:
        row.cells[0].merge(row.cells[num_cols - 1])
    cell = row.cells[0]
    _set_cell_shading(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = FONT_NAME
    r.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_margins(cell)
    return row


def _make_col_header_row(table, headers, col_widths=None):
    row = table.add_row()
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _set_cell_shading(cell, "DAE9F7")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = FONT_NAME
        r.font.color.rgb = RGBColor(0, 51, 102)
        _set_cell_margins(cell)
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])
    return row


def _formatear_numero(valor: float | int) -> str:
    entero = int(valor)
    decimales = int(round((valor - entero) * 100))
    s = f"{entero:,}".replace(",", ".")
    return f"${s}.{decimales:02d}"


def generar_docx(contratista: dict, contratos: list, resumen: dict) -> bytes:
    """Genera DOCX con estructura INFORME DE ACTIVIDADES.

    Formato basado en el documento de referencia:
      - Título centrado: INFORME DE ACTIVIDADES No. XX-2026
      - Párrafo introductorio
      - Tabla IDENTIFICACIÓN CONTRACTUAL
      - Tabla ACTIVIDADES GENERALES (4 cols)
      - Tabla ACTIVIDADES ESPECÍFICAS (4 cols)
      - Firma del contratista
      - ANEXOS
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    today = datetime.now()
    c = contratos[0] if contratos else {}
    perfil = c.get("perfil", "") or ""

    # ─── TITLE ───────────────────────────────────────────────────────
    _add_paragraph(doc, f"INFORME DE ACTIVIDADES No. {c.get('numero_contrato', '')}-{today.year}",
                   bold=True, size=FONT_SIZE_TITLE, color=COLOR_PRIMARY,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ─── INTRO ───────────────────────────────────────────────────────
    periodo = ""
    if c.get("fecha_inicio") and c.get("fecha_fin"):
        try:
            fi = datetime.strptime(c["fecha_inicio"], "%Y-%m-%d")
            ff = datetime.strptime(c["fecha_fin"], "%Y-%m-%d")
            periodo = f"Del {fi.day:02d} de {MESES_TITULO[fi.month]} al {ff.day:02d} de {MESES_TITULO[ff.month]} del {ff.year}"
        except (ValueError, TypeError):
            periodo = f"{MESES_TITULO[today.month]} de {today.year}"
    else:
        periodo = f"{MESES_TITULO[today.month]} de {today.year}"

    _add_paragraph(doc,
        f"De acuerdo con el Contrato suscrito con la Empresa Social del Estado Norte 3 – E.S.E., "
        f"me permito presentar el informe de actividades ejecutadas durante el periodo mencionado, "
        f"con el fin de acreditar su cumplimiento:",
        size=FONT_SIZE, space_after=12)

    # ─── TABLE 1: IDENTIFICACIÓN CONTRACTUAL ─────────────────────────
    t1 = doc.add_table(rows=1, cols=2)
    _set_table_borders(t1, sz="6", color="003366")

    # Header
    _make_header_row(t1, "IDENTIFICACIÓN CONTRACTUAL", 2, "003366")

    # Identify fields based on available data
    tipo_doc = "CC" if contratista.get("tipo_persona") == "NATURAL" else "NIT"
    obj = c.get("objeto", "")
    obj_short = obj[:120] + "..." if len(obj) > 120 else obj

    # Determine periodo from contract or pago data
    periodo_ejecutado = periodo
    if c.get("periodo_desde") and c.get("periodo_hasta"):
        periodo_ejecutado = f"{c['periodo_desde']} al {c['periodo_hasta']}"

    # Valor en letras
    monto = float(c.get("monto_total", 0) or 0)
    valor_final = float(c.get("valor_final", 0) or monto)
    valor_letras = c.get("valor_letras", "") or ""

    supervisor_line = c.get("supervisor", "")
    if c.get("cargo_supervisor"):
        supervisor_line += f" {c['cargo_supervisor']}"
    if c.get("unidad_atencion"):
        supervisor_line += f" de {c['unidad_atencion']}"

    ident_rows = [
        ("NÚMERO Y FECHA DE CONTRATO", f"{c.get('numero_contrato', '')}{' del ' + c.get('fecha_contrato', '') if c.get('fecha_contrato') else ''}"),
        ("TIPO DE CONTRATO", "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES"),
        ("CONTRATANTE", "EMPRESA SOCIAL DEL ESTADO NORTE 3 – ESE\nNIT. 900.146.438.-4"),
        ("CONTRATISTA", contratista.get("nombre", "")),
        ("IDENTIFICACIÓN", f"{tipo_doc}. {contratista.get('identificacion', '')}"),
        ("OBJETO DEL CONTRATO", obj_short),
        ("VALOR DEL CONTRATO", f"{valor_letras} ({_formatear_numero(valor_final)})"),
        ("TÉRMINO DEL CONTRATO", f"Desde la fecha de suscripción del acta de inicio, previo registro presupuestal y aprobación de garantías, hasta el {c.get('fecha_fin', '')} previo cumplimiento de los requisitos de perfeccionamiento y ejecución del contrato."),
        ("PERIODO EJECUTADO", periodo_ejecutado),
        ("SUPERVISOR DESIGNADO", supervisor_line),
    ]

    for i, (label, value) in enumerate(ident_rows):
        bg = "F0F4FA" if i % 2 == 0 else "FFFFFF"
        row = t1.add_row()
        _set_cell_shading(row.cells[0], bg)
        _add_cell_text(row.cells[0], label, bold=True, size=9, color=COLOR_TEXT,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_margins(row.cells[0])
        row.cells[0].width = Cm(4.5)

        _set_cell_shading(row.cells[1], bg)
        _add_cell_text(row.cells[1], value, bold=False, size=9, color=COLOR_TEXT,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if len(value) > 60 else WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_margins(row.cells[1])
        row.cells[1].width = Cm(12.0)

    # Apply cell margins to header too
    for cell in t1.rows[0].cells:
        _set_cell_margins(cell)
    _set_cell_shading(t1.rows[0].cells[0], "003366")

    _add_paragraph(doc, "", size=6, space_after=6)  # spacer

    # ─── SPLIT ACTIVITIES: GENERALES vs ESPECÍFICAS ──────────────────
    acts_generales = [a for a in c.get("actividades", []) if a.get("tipo") == "GENERAL" or a.get("tipo") is None]
    acts_especificas = [a for a in c.get("actividades", []) if a.get("tipo") == "ESPECIFICA"]

    # If all activities have the same type or no type, put all as generales
    if not acts_especificas:
        acts_especificas = []
    if not acts_generales and not acts_especificas:
        acts_generales = c.get("actividades", [])

    def _build_activity_table(doc, title, actividades, col_widths=[1.0, 5.5, 6.0, 4.0]):
        if not actividades:
            return

        _add_paragraph(doc, title, bold=True, size=11, color=COLOR_PRIMARY,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

        tbl = doc.add_table(rows=1, cols=4)
        _set_table_borders(tbl, sz="6", color="003366")

        # Merged header
        _make_header_row(tbl, title, 4, "003366")

        # Column headers
        _make_col_header_row(tbl, ["No.", "ACTIVIDAD CONTRATADA", "DESCRIPCIÓN DE ACCIONES REALIZADAS", "EVIDENCIAS"], col_widths)

        # Data rows
        for i, act in enumerate(actividades):
            num = i + 1
            desc = act.get("descripcion", "")
            evidencias = act.get("evidencias", [])

            # Build description from evidencias (TEXTO type)
            acciones = desc if desc else ""
            textos = [e.get("contenido_texto", "") for e in evidencias if e.get("tipo") == "TEXTO" and e.get("contenido_texto")]
            if textos and not acciones:
                acciones = "\n".join(textos)

            # Build evidence text
            ev_text_parts = []
            for ev in evidencias:
                if ev.get("tipo") == "IMAGEN":
                    ev_text_parts.append(ev.get("archivo_nombre", "Imagen"))
                elif ev.get("tipo") == "ARCHIVO":
                    ev_text_parts.append(ev.get("archivo_nombre", "Archivo"))
                elif ev.get("tipo") == "TEXTO":
                    txt = ev.get("contenido_texto", "")
                    if txt and not acciones:
                        # Already used as description
                        pass
            ev_text = "\n".join(ev_text_parts) if ev_text_parts else ""

            alt_bg = "F5F8FC" if i % 2 == 0 else "FFFFFF"
            row = tbl.add_row()
            vals = [str(num), desc, acciones, ev_text]
            for ci, val in enumerate(vals):
                cell = row.cells[ci]
                _set_cell_shading(cell, alt_bg)
                _set_cell_margins(cell)
                if ci < len(col_widths):
                    cell.width = Cm(col_widths[ci])
                if ci in (1, 2) and val:
                    # Columnas con HTML: renderizar con formato real
                    _render_html_to_cell(cell, val, font_size=9)
                else:
                    # Columnas de texto plano (No., Evidencias)
                    align = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
                    _add_cell_text(cell, val, bold=(ci == 0), size=9,
                                   color=COLOR_TEXT, alignment=align)

    # ─── TABLE 2: ACTIVIDADES GENERALES ──────────────────────────────
    _build_activity_table(doc, "ACTIVIDADES GENERALES", acts_generales)

    # ─── TABLE 3: ACTIVIDADES ESPECÍFICAS ────────────────────────────
    _build_activity_table(doc, "ACTIVIDADES ESPECÍFICAS", acts_especificas)

    # ─── CLOSING ─────────────────────────────────────────────────────
    _add_paragraph(doc, "", size=6, space_after=12)  # spacer

    _add_paragraph(doc, "Cordialmente,", bold=False, size=FONT_SIZE, space_before=12, space_after=18)

    _add_paragraph(doc, contratista.get("nombre", ""),
                   bold=True, size=FONT_SIZE, color=COLOR_PRIMARY, space_after=2)
    _add_paragraph(doc, f"CC. {contratista.get('identificacion', '')}",
                   bold=True, size=FONT_SIZE, color=COLOR_PRIMARY, space_after=2)
    _add_paragraph(doc, perfil,
                   bold=False, size=FONT_SIZE, space_after=2)
    _add_paragraph(doc, "Contratista",
                   bold=False, size=FONT_SIZE, space_after=12)

    # ─── ANEXOS ──────────────────────────────────────────────────────
    _add_paragraph(doc, "ANEXOS",
                   bold=True, size=11, color=COLOR_PRIMARY,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=6)

    # ─── OUTPUT ──────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
