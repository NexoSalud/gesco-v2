"""Generador de documentos .docx para actas de entrega y devolución de inventario."""

import io
import os
import re
from datetime import date
from docx import Document
from docx.shared import Pt

MESES_ESPANOL = [
    "", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]

TEMPLATE_FILENAMES = {
    ("ENTREGA", "TECNOLOGICO"): "ACTA ENTREGA TECNOLOGICO.docx",
    ("DEVOLUCION", "TECNOLOGICO"): "ACTA DEVOLUCION TECNOLOGICO.docx",
    ("ENTREGA", "BIOMEDICO"): "ACTA ENTREGA BIOMEDICO.docx",
    ("DEVOLUCION", "BIOMEDICO"): "ACTA DEVOLUCION BIOMEDICO.docx",
    ("ENTREGA", "DOTACION"): "ACTA ENTREGA DOTACION.docx",
    ("DEVOLUCION", "DOTACION"): "ACTA DEVOLUCION DOTACION.docx",
    ("ENTREGA", "DOTACION_INSUMO"): "ACTA ENTREGA DOTACION.docx",
    ("DEVOLUCION", "DOTACION_INSUMO"): "ACTA DEVOLUCION DOTACION.docx",
    ("ENTREGA", "INSUMO"): "ACTA ENTREGA DOTACION.docx",
    ("DEVOLUCION", "INSUMO"): "ACTA DEVOLUCION DOTACION.docx",
}


def _replace_text_in_paragraph(paragraph, day, month, year, contratista_name, cc, resolution, cargo):
    original_text = paragraph.text
    if not original_text.strip():
        return False
        
    replaced_text = original_text
    
    # 1. CC: "cédula de ciudadanía No. X+"
    replaced_text = re.sub(r'(cédula de ciudadanía No\.\s*)[Xx]+', rf'\g<1>{cc}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(cédula de ciudadanía No\s*)[Xx]+', rf'\g<1>{cc}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(cédula de ciudadanía de No\.\s*)[Xx]+', rf'\g<1>{cc}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'No\.\s*[Xx]+(?=\s*,?\s*recibe)', f'No. {cc}', replaced_text, flags=re.IGNORECASE)
    
    # 2. Day: "a los X+ días"
    replaced_text = re.sub(r'(a los\s+)[Xx]+(\s+días)', rf'\g<1>{day}\g<2>', replaced_text, flags=re.IGNORECASE)
    
    # 3. Month: "del mes de X+"
    replaced_text = re.sub(r'(del mes de\s+)[Xx]+', rf'\g<1>{month}', replaced_text, flags=re.IGNORECASE)
    
    # 4. Year: "de X+ se suscribe"
    replaced_text = re.sub(r'(de\s+)[Xx]+(?=\s+,?\s*se suscribe)', rf'\g<1>{year}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(julio de\s+)[Xx]+(?=\s+,?\s*se suscribe)', rf'\g<1>{year}', replaced_text, flags=re.IGNORECASE)
    
    # 5. Contratista in main text: "contratista X+"
    replaced_text = re.sub(r'(/ contratista\s+)[Xx]+', rf'\g<1>{contratista_name}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(contratista\s+)[Xx]+', rf'\g<1>{contratista_name}', replaced_text, flags=re.IGNORECASE)
    
    # 6. Resolution: "Resolución No. X+"
    replaced_text = re.sub(r'(Resolución No\.\s*)[Xx]+', rf'\g<1>{resolution}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(Resolución\s*)[Xx]+', rf'\g<1>{resolution}', replaced_text, flags=re.IGNORECASE)
    
    # 7. Signature blocks at bottom
    replaced_text = re.sub(r'(NOMBRES:\s*)[Xx]+', rf'\g<1>{contratista_name}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(CÉDULA:\s*)[Xx]+', rf'\g<1>{cc}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(CEDULA:\s*)[Xx]+', rf'\g<1>{cc}', replaced_text, flags=re.IGNORECASE)
    replaced_text = re.sub(r'(CARGO:\s*)[Xx]+', rf'\g<1>{cargo}', replaced_text, flags=re.IGNORECASE)
    
    if replaced_text != original_text:
        # Save formatting: clear runs and write to first run
        for i, run in enumerate(paragraph.runs):
            run.text = replaced_text if i == 0 else ""
        return True
    return False


def _replace_all_text(doc, day, month, year, contratista_name, cc, resolution, cargo):
    # Paragraphs
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, day, month, year, contratista_name, cc, resolution, cargo)
    # Table cells (excluding the first table which contains the items list)
    for table in doc.tables[1:]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, day, month, year, contratista_name, cc, resolution, cargo)
                    
    # Also replace in the signature/footer section if there are tables
    if len(doc.tables) > 0:
        # Check the last table specifically, as it might contain signature blocks
        last_table = doc.tables[-1]
        for row in last_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, day, month, year, contratista_name, cc, resolution, cargo)


def _populate_items_table(table, category, items):
    # 1. Clear all rows except the first (header)
    while len(table.rows) > 1:
        tbl = table._tbl
        tr = table.rows[1]._tr
        tbl.remove(tr)
        
    # 2. Add rows for each item
    for idx, item in enumerate(items, start=1):
        row = table.add_row()
        
        # Determine cell values based on category
        if category == "TECNOLOGICO":
            vals = [
                str(idx),
                str(item.get("tipo_elemento") or ""),
                str(item.get("marca") or ""),
                str(item.get("modelo") or ""),
                str(item.get("serial") or ""),
                str(item.get("imei2") or "N/A"),
                str(item.get("estado_declarado") or "Excelente"),
                str(item.get("observaciones") or "")
            ]
        elif category == "BIOMEDICO":
            vals = [
                str(idx),
                str(item.get("tipo_elemento") or ""),
                str(item.get("marca") or ""),
                str(item.get("modelo") or ""),
                str(item.get("serial") or ""),
                str(item.get("estado_declarado") or "Excelente"),
                str(item.get("observaciones") or "")
            ]
        else:  # DOTACION
            vals = [
                str(idx),
                str(item.get("tipo_elemento") or ""),
                str(item.get("cantidad") or "1"),
                str(item.get("estado_declarado") or "Excelente"),
                str(item.get("observaciones") or "")
            ]
            
        # Fill cells and style
        for col_idx, val in enumerate(vals):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell.text = val
                for paragraph in cell.paragraphs:
                    # Apply styling (9pt font size)
                    for run in paragraph.runs:
                        run.font.name = 'Aptos Display'
                        run.font.size = Pt(9)


def generar_acta_inventario_docx(tipo: str, categoria: str, contrato_data: dict, items: list, fecha_acta: date = None) -> bytes:
    """Genera el .docx del acta reemplazando encabezados, firmantes y agregando los items a la tabla.

    Args:
        tipo: 'ENTREGA' o 'DEVOLUCION'
        categoria: 'TECNOLOGICO', 'BIOMEDICO', o 'DOTACION'
        contrato_data: dict con info de contratista, cedula, resolucion, cargo
        items: list of dicts con las columnas a insertar en la tabla
        fecha_acta: date opcional, hoy por defecto
    """
    if not fecha_acta:
        fecha_acta = date.today()

    template_name = TEMPLATE_FILENAMES.get((tipo, categoria))
    if not template_name:
        raise ValueError(f"No existe plantilla para el tipo={tipo} y categoría={categoria}")

    template_path = os.path.normpath(os.path.join(
        os.path.dirname(__file__), "..", "templates", template_name))

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla de acta no encontrada: {template_path}")

    doc = Document(template_path)

    # Preparar campos de texto
    day = str(fecha_acta.day)
    month = MESES_ESPANOL[fecha_acta.month]
    year = str(fecha_acta.year)
    contratista_name = (contrato_data.get("nombre_contratista") or "").upper()
    cc = contrato_data.get("cedula") or ""
    resolution = contrato_data.get("resolucion_codigo") or ""
    cargo = (contrato_data.get("perfil") or "").upper()

    # 1. Reemplazar encabezados y firmas
    _replace_all_text(doc, day, month, year, contratista_name, cc, resolution, cargo)

    # 2. Rellenar la tabla de elementos (siempre es la primera tabla del documento)
    if len(doc.tables) > 0:
        _populate_items_table(doc.tables[0], categoria, items)

    # Guardar a bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
