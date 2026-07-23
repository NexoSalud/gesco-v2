"""Generador de .docx de contratos. Carga plantilla, reemplaza placeholders e inserta
obligaciones con soporte de tablas HTML renderizadas como tablas DOCX reales."""
import io, os, re
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.services.numero_letras import numero_a_letras

MESES_ESPANOL = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _formatear_numero(valor: float | int) -> str:
    """Formatea un número al estilo colombiano: $20.000.000.00
    Usa punto como separador de miles y dos decimales."""
    entero = int(valor)
    decimales = int(round((valor - entero) * 100))
    s = f"{entero:,}".replace(",", ".")
    return f"${s}.{decimales:02d}"


def formatear_fecha(valor: str | date | None, default: str = "_________") -> str:
    """Convierte fecha a formato español: '06 de julio de 2026'.
    Acepta string ISO, date object, o None."""
    if not valor:
        return default
    if isinstance(valor, str):
        try:
            partes = valor.split("-")
            if len(partes) == 3:
                d = date(int(partes[0]), int(partes[1]), int(partes[2]))
            else:
                return valor
        except (ValueError, IndexError):
            return valor
    elif isinstance(valor, date):
        d = valor
    else:
        return str(valor)
    return f"{d.day:02d} de {MESES_ESPANOL[d.month].lower()} de {d.year}"


TEMPLATE_PATH = os.path.normpath(os.path.join(
    os.path.dirname(__file__), "..", "templates", "plantilla_contrato.docx"))


def _merge_runs_and_replace(paragraph, old_text, new_text):
    if old_text not in paragraph.text:
        return False
    full_new = paragraph.text.replace(old_text, str(new_text) if new_text else "")
    for i, run in enumerate(paragraph.runs):
        run.text = full_new if i == 0 else ""
    return True


def _replace_all(doc, placeholders):
    for ph, val in placeholders.items():
        val = str(val) if val is not None else ""
        for p in doc.paragraphs:
            if ph in p.text:
                _merge_runs_and_replace(p, ph, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if ph in p.text:
                            _merge_runs_and_replace(p, ph, val)


def _parse_html_blocks(html_text):
    """Retorna [(tipo, contenido)] donde tipo='text'|'table'."""
    blocks, pos = [], 0
    while pos < len(html_text):
        # Buscar <table> usando el metodo mas simple: buscar por substring
        table_start = html_text.lower().find('<table', pos)
        if table_start == -1:
            rest = html_text[pos:].strip()
            if rest:
                blocks.append(('text', rest))
            break
        # Texto antes de la tabla
        if table_start > pos and html_text[pos:table_start].strip():
            blocks.append(('text', html_text[pos:table_start].strip()))
        # Encontrar </table>
        table_end = html_text.lower().find('</table>', table_start + 6)
        if table_end == -1:
            # No se encontro cierre, tratar como texto
            blocks.append(('text', html_text[pos:].strip()))
            break
        table_end += 8  # len('</table>')
        blocks.append(('table', html_text[table_start:table_end]))
        pos = table_end
    return blocks


def _extract_table_data(table_html):
    """Extrae (headers, rows) de un HTML <table>."""
    rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
    headers, rows = [], []
    for rh in rows_html:
        cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', rh, re.DOTALL | re.IGNORECASE)
        texts = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
        if '<th' in rh.lower():
            headers = texts
        else:
            rows.append(texts)
    return headers, rows


def _make_table_xml(headers, rows_data):
    """Crea elemento <w:tbl> con bordes."""
    ns = 'w'
    tbl = OxmlElement(f'{ns}:tbl')
    tblPr = OxmlElement(f'{ns}:tblPr')
    tw = OxmlElement(f'{ns}:tblW'); tw.set(qn('w:w'), '5000'); tw.set(qn('w:type'), 'pct')
    tblPr.append(tw)
    tb = OxmlElement(f'{ns}:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'{ns}:{side}'); b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000')
        tb.append(b)
    tblPr.append(tb)
    tbl.append(tblPr)
    cols = max(len(headers) if headers else 1, *(len(r) for r in rows_data), 1)
    tg = OxmlElement(f'{ns}:tblGrid')
    for _ in range(cols):
        gc = OxmlElement(f'{ns}:gridCol'); gc.set(qn('w:w'), str(int(5000/cols)))
        tg.append(gc)
    tbl.append(tg)

    def _cell(text, bold=False):
        tc = OxmlElement(f'{ns}:tc')
        tcW = OxmlElement(f'{ns}:tcW'); tcW.set(qn('w:w'),str(int(5000/cols))); tcW.set(qn('w:type'),'pct')
        tc.append(tcW)
        if bold:
            shd = OxmlElement(f'{ns}:shd'); shd.set(qn('w:val'),'clear')
            shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
            tc.append(shd)
        p = OxmlElement(f'{ns}:p')
        r = OxmlElement(f'{ns}:r')
        rPr = OxmlElement(f'{ns}:rPr')
        rf = OxmlElement(f'{ns}:rFonts'); rf.set(qn('w:ascii'),'Aptos Display'); rf.set(qn('w:hAnsi'),'Aptos Display')
        rPr.append(rf)
        sz = OxmlElement(f'{ns}:sz'); sz.set(qn('w:val'),'20')  # 10pt para tablas
        rPr.append(sz)
        if bold:
            b = OxmlElement(f'{ns}:b'); rPr.append(b)
        r.append(rPr)
        t = OxmlElement(f'{ns}:t'); t.text=str(text).strip(); t.set(qn('xml:space'),'preserve')
        r.append(t); p.append(r); tc.append(p)
        return tc

    if headers:
        tr = OxmlElement(f'{ns}:tr')
        for h in headers: tr.append(_cell(h, bold=True))
        tbl.append(tr)
    for row in rows_data:
        tr = OxmlElement(f'{ns}:tr')
        for cv in row: tr.append(_cell(cv))
        while len(list(tr)) < cols: tr.append(_cell(''))
        tbl.append(tr)
    return tbl


def _make_paragraph_xml(text, bold=False, size=11):
    """Crea elemento <w:p> con formato 11pt, justify, Aptos Display."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'both'); pPr.append(jc)
    if bold:
        pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'),'Normal')
        pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Aptos Display'); rf.set(qn('w:hAnsi'),'Aptos Display')
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'22'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'),'22'); rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text=text; t.set(qn('xml:space'),'preserve')
    r.append(t); p.append(r)
    return p


def generar_contrato_docx(data: dict, obligaciones_esp: list[str] | None = None) -> bytes:
    valor = float(data.get("valor_contrato", 0))
    valor_letras = data.get("valor_letras", "") or numero_a_letras(valor)

    placeholders = {
        "<<NO. DE CONTRATO>>": data.get("numero_contrato", "_________"),
        "<<FECHA DEL CONTRATO>>": formatear_fecha(data.get("fecha_contrato", data.get("fecha_inicio", "_________"))),
        "<<fecha del contrato>>": formatear_fecha(data.get("fecha_contrato", data.get("fecha_inicio", "_________"))),
        "<<CONTRATISTA>>": data.get("nombre_contratista", "___________________"),
        "<<CEDULA DEL CONTRATISTA>>": data.get("cedula", "_________"),
        "<<CÉDULA DEL CONTRATISTA>>": data.get("cedula", "_________"),
        "<<LUGAR DE EXPEDICIÓN>>": data.get("lugar_expedicion", "_________"),
        "<< LUGAR DE EXPEDICIÓN>>": data.get("lugar_expedicion", "_________"),
        "<<DIRECCIÓN>>": data.get("direccion", "_________"),
        "<<TELÉFONO>>": data.get("telefono", "_________"),
        "<<CORREO>>": data.get("correo", "_________"),
        "<<OBJETO DEL CONTRATO>>": data.get("objeto", "_________"),
        "<<FECHA DE TERMINACIÓN>>": formatear_fecha(data.get("fecha_fin", "_________")),
        "<<fecha de terminación>>": formatear_fecha(data.get("fecha_fin", "_________")),
        "<<VALOR DEL CONTRATO>>": f"{valor_letras} ({_formatear_numero(valor)})",
        "<<CDP>>": data.get("no_cdp", "_________"),
        "<<FECHA DEL CDP>>": formatear_fecha(data.get("fecha_cdp", "") or ""),
        "<<fecha del CDP>>": formatear_fecha(data.get("fecha_cdp", "") or ""),
        "<<VALOR DEL CDP>>": _formatear_numero(valor),
        "<<LUGAR DE EJECUCIÓN>>": data.get("lugar_ejecucion", "Puerto Tejada - Cauca"),
        "<<fecha del acta>>": formatear_fecha(data.get("fecha_inicio", str(date.today()))),
        "<<SUPERVISOR>>": data.get("supervisor", "_________"),
        "<<CEDULA DE SUPERVISOR>>": data.get("cedula_supervisor", "_________"),
        "<<CÉDULA DE SUPERVISOR>>": data.get("cedula_supervisor", "_________"),
        "<<PERFIL>>": data.get("perfil", "_________"),
        "<<RUBRO>>": data.get("rubro", "") or data.get("imputacion", "") or "_________",
    }

    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
            sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)

    _replace_all(doc, placeholders)

    # Insertar obligaciones con soporte de tablas HTML
    if obligaciones_esp:
        for pi, p in enumerate(doc.paragraphs):
            if "<<OBLIGACIONES>>" in p.text:
                _merge_runs_and_replace(p, "<<OBLIGACIONES>>", "")
                last_element = p._p
                for oblig in obligaciones_esp:
                    es_header = oblig.strip().upper() in ('OBLIGACIONES GENERALES:', 'OBLIGACIONES ESPECÍFICAS:') or oblig.strip().startswith('OBLIGACIONES GENERALES') or oblig.strip().startswith('OBLIGACIONES ESPECÍFICAS')
                    blocks = _parse_html_blocks(oblig)
                    for bi, (btype, content) in enumerate(blocks):
                        if btype == 'table':
                            headers, rows = _extract_table_data(content)
                            if headers or rows:
                                tbl = _make_table_xml(headers, rows)
                                last_element.addnext(tbl)
                                last_element = tbl
                        else:
                            clean = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
                            clean = re.sub(r'<[^>]+>', '', clean)
                            clean = clean.replace('&nbsp;',' ').replace('&amp;','&')
                            clean = re.sub(r'\n{3,}', '\n\n', clean).strip()
                            parts = clean.split('\n')
                            for pi2, part in enumerate(parts):
                                part = part.strip()
                                if not part:
                                    continue
                                new_p = _make_paragraph_xml(part, bold=False, size=11)
                                last_element.addnext(new_p)
                                last_element = new_p
                break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── Documentos de contratación (INEXISTENCIA, ESTUDIOS PREVIOS, etc.) ────────

DOCUMENTOS_TEMPLATES = {
    "inexistencia": "inexistencia.docx",
    "estudios_previos": "estudios_previos.docx",
    "solicitud_cdp": "solicitud_cdp.docx", 
    "invitacion": "invitacion.docx",
    "idoneidad": "idoneidad.docx",
    "designacion_supervision": "designacion_supervision.docx",
    "acta_inicio": "acta_inicio.docx",
    "acta_liquidacion": "acta_liquidacion.docx",
}


def generar_documento_contrato(tipo: str, data: dict) -> bytes:
    """Genera un documento de contratacion (inexistencia, estudios previos, etc.)
    a partir de su plantilla DOCX y los datos del contrato.

    Args:
        tipo: uno de 'inexistencia', 'estudios_previos', 'solicitud_cdp',
              'invitacion', 'idoneidad'
        data: dict con los datos del contrato

    Returns:
        bytes del archivo .docx
    """
    from datetime import date

    template_name = DOCUMENTOS_TEMPLATES.get(tipo)
    if not template_name:
        raise ValueError(f"Tipo de documento desconocido: {tipo}")

    template_path = os.path.normpath(os.path.join(
        os.path.dirname(__file__), "..", "templates", template_name))

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")

    doc = Document(template_path)

    # Mapa de placeholders comun para todos los documentos
    valor = float(data.get("monto_total", 0))
    from app.services.numero_letras import numero_a_letras
    valor_letras = data.get("valor_letras", "") or numero_a_letras(valor)

    hoy = date.today()

    placeholders = {
        "<<NO. DE CONTRATO>>": data.get("numero_contrato", "_________"),
        "<<No. DE CONTRATO>>": data.get("numero_contrato", "_________"),
        "<<FECHA DEL CONTRATO>>": formatear_fecha(data.get("fecha_contrato", data.get("fecha_inicio", str(hoy)))),
        "<<fecha del contrato>>": formatear_fecha(data.get("fecha_contrato", data.get("fecha_inicio", str(hoy)))),
        "<<CONTRATISTA>>": data.get("nombre_contratista", "___________________"),
        "<<CEDULA DEL CONTRATISTA>>": data.get("cedula", "_________"),
        "<<CÉDULA DEL CONTRATISTA>>": data.get("cedula", "_________"),
        "<<cédula del contratista>>": data.get("cedula", "_________"),
        "<<cédula del supervisor>>": data.get("cedula_supervisor", "_________"),
        "<<SUPERVISOR>>": data.get("supervisor", "_________"),
        "<<CEDULA DE SUPERVISOR>>": data.get("cedula_supervisor", "_________"),
        "<<CÉDULA DE SUPERVISOR>>": data.get("cedula_supervisor", "_________"),
        "<<OBJETO DEL CONTRATO>>": data.get("objeto", "_________"),
        "<<PERFIL>>": data.get("perfil", "_________"),
        "<<VALOR DEL CONTRATO>>": f"{valor_letras} ({_formatear_numero(valor)})",
        "<<VALOR DE CDP>>": _formatear_numero(valor).lstrip("$"),
        "<<fecha de terminación>>": formatear_fecha(data.get("fecha_fin", "_________")),
        "<<fecha de finalización>>": formatear_fecha(data.get("fecha_fin", "_________")),
        "<<FECHA DE TERMINACIÓN>>": formatear_fecha(data.get("fecha_fin", "_________")),
        "<<día>>": str(hoy.day),
        "<<mes>>": MESES_ESPANOL[hoy.month].upper(),
        "<<año>>": str(hoy.year),
        "<<Unidad de Atención>>": data.get("unidad_atencion", "_________"),
        "<<unidad de atención>>": data.get("unidad_atencion", "_________"),
        "<<LUGAR DE EXPEDICIÓN>>": data.get("lugar_expedicion", "_________"),
        "<< LUGAR DE EXPEDICIÓN>>": data.get("lugar_expedicion", "_________"),
        "<<lugar de expedición>>": data.get("lugar_expedicion", "_________"),
        "<<DIRECCIÓN>>": data.get("direccion", "_________"),
        "<<TELÉFONO>>": data.get("telefono", "_________"),
        "<<CORREO>>": data.get("correo", "_________"),
        "<<CDP>>": data.get("no_cdp", "_________"),
        "<<FECHA DEL CDP>>": formatear_fecha(data.get("fecha_cdp", "") or ""),
        "<<fecha del CDP>>": formatear_fecha(data.get("fecha_cdp", "") or ""),
        "<<VALOR DEL CDP>>": _formatear_numero(valor),
        "<<UNSPSC>>": data.get("codigo_unspsc", ""),
        "<<DESCRIPCIÓN>>": data.get("descripcion_unspsc", ""),
        "<<LUGAR DE EJECUCIÓN>>": data.get("lugar_ejecucion", "Puerto Tejada - Cauca"),
        "<<OBLIGACIONES>>": data.get("obligaciones", "Ver cláusula SEGUNDA del contrato."),
        "<<RUBRO>>": data.get("rubro", "") or data.get("imputacion", "") or "_________",
    }

    # Reemplazar placeholders (recursivo para tablas anidadas)
    def _replace_in_element(element, ph, val):
        """Busca y reemplaza un placeholder en paragraphs y tablas anidadas."""
        for p in element.paragraphs:
            if ph in p.text:
                full_new = p.text.replace(ph, val)
                for i, run in enumerate(p.runs):
                    run.text = full_new if i == 0 else ""
        # Buscar en tablas anidadas (recursivo)
        from docx.oxml.ns import qn
        for tbl_elem in element._element.findall(qn('w:tbl')):
            from docx.table import Table
            nested_table = Table(tbl_elem, element)
            for row in nested_table.rows:
                for cell in row.cells:
                    _replace_in_element(cell, ph, val)

    import re as _re
    # ─── Helper: convertir HTML a texto plano legible ────────────────────
    def _html_a_texto(html_content: str) -> str:
        """Convierte HTML a texto plano, preservando saltos de línea y viñetas."""
        if not html_content or ('<' not in html_content or '>' not in html_content):
            return html_content
        texto = html_content
        # Reemplazar <br> con nueva línea
        texto = re.sub(r'<br\s*/?>', '\n', texto, flags=re.IGNORECASE)
        # </p> → nueva línea
        texto = re.sub(r'</p>', '\n', texto, flags=re.IGNORECASE)
        # <li> → viñeta
        texto = re.sub(r'<li[^>]*>', '• ', texto, flags=re.IGNORECASE)
        texto = re.sub(r'</li>', '\n', texto, flags=re.IGNORECASE)
        # <tr> → nueva línea
        texto = re.sub(r'</tr>', '\n', texto, flags=re.IGNORECASE)
        # <td>, <th> → separar con tabulación
        texto = re.sub(r'</t[dh]>', '\t', texto, flags=re.IGNORECASE)
        # Limpiar todos los demás tags HTML
        texto = re.sub(r'<[^>]+>', '', texto)
        # Entidades HTML
        texto = texto.replace('&nbsp;', ' ').replace('&amp;', '&')
        texto = texto.replace('&lt;', '<').replace('&gt;', '>')
        # Limpiar líneas en blanco excesivas
        texto = re.sub(r'\n{3,}', '\n\n', texto)
        texto = re.sub(r'\t{2,}', '\t', texto)
        return texto.strip()

    # ─── Convertir valores HTML a texto plano ────────────────────────────
    oblig_plano = _html_a_texto(placeholders.get("<<OBLIGACIONES>>", ""))
    objeto_plano = _html_a_texto(placeholders.get("<<OBJETO DEL CONTRATO>>", ""))

    # Placeholders: reemplazo simple de texto (body + tablas)
    for ph, val in placeholders.items():
        val = str(val) if val is not None else ""
        # Usar valor limpio (sin HTML) para OBLIGACIONES y OBJETO
        if ph == "<<OBLIGACIONES>>":
            val_clean = oblig_plano if oblig_plano else "Ver cláusula SEGUNDA del contrato."
        elif ph == "<<OBJETO DEL CONTRATO>>":
            val_clean = objeto_plano
        else:
            val_clean = val
        # 1. Body paragraphs
        for p in doc.paragraphs:
            if ph in p.text:
                full_new = p.text.replace(ph, val_clean)
                for i, run in enumerate(p.runs):
                    run.text = full_new if i == 0 else ""
        # 2. Top-level tables (con recursión para anidadas)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_in_element(cell, ph, val_clean)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
