"""Generador de Informe Mensual de Actividades — Apoyo Administrativo EBS.

Genera un documento DOCX con estructura idéntica al PDF de referencia:
  - Carta al Sindicato
  - I. Objeto contractual y alcance
  - II. Metodología de ejecución
  - III. Relación de perfiles
  - IV. Tabla de afiliados partícipes
  - V. Actividades generales
  - VI. Actividades específicas por perfil funcional
"""

import io
import os
import logging
from datetime import date, datetime as dt
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from sqlalchemy import select, func as sa_func
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.apoyo_administrativo import ApoyoAdministrativo
from app.models.actividad_apoyo import ActividadApoyo
from app.models.evidencia_apoyo import EvidenciaApoyo

logger = logging.getLogger(__name__)

MESES = [
    "", "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
    "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE",
]

MESES_TITULO = [
    "", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]

# ─── Constantes de estilo ─────────────────────────────────────────────────

FONT_NAME = "Calibri"
FONT_SIZE_NORMAL = Pt(10)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_TITLE = Pt(12)
FONT_SIZE_SECTION = Pt(11)
FONT_SIZE_TABLE_HEADER = Pt(8)
FONT_SIZE_TABLE_CELL = Pt(8)

COLOR_PRIMARY = RGBColor(0x00, 0x3F, 0x72)      # Azul oscuro corporativo
COLOR_SECONDARY = RGBColor(0x00, 0x5B, 0x96)     # Azul medio
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)           # Texto oscuro
COLOR_TABLE_HEADER_BG = "003F72"                  # Fondo azul oscuro
COLOR_TABLE_LIGHT_BG = "EBF5FB"                   # Fondo azul claro

MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.5)


def _set_cell_shading(cell, color: str):
    """Aplica color de fondo a una celda."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda. kwargs: top, bottom, left, right = size."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, sz in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(sz))
        element.set(qn("w:color"), "000000")
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_formatted_paragraph(doc_or_cell, text: str, bold=False, italic=False,
                              size=FONT_SIZE_NORMAL, color=COLOR_TEXT,
                              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              space_after=Cm(0.2), space_before=Cm(0),
                              font_name=FONT_NAME, line_spacing=1.15):
    """Agrega un párrafo con formato."""
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def _add_table_header_row(table, headers: list[str], col_widths: list[float] = None):
    """Agrega fila de encabezado a una tabla con formato."""
    hdr = table.rows[0] if table.rows else table.add_row()
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        _set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TABLE_HEADER
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])


def _add_data_row(table, values: list[str], col_widths: list[float] = None,
                   bold_first=False, bg_color=None):
    """Agrega una fila de datos a una tabla."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        if bg_color:
            _set_cell_shading(cell, bg_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(val)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TABLE_CELL
        run.font.color.rgb = COLOR_TEXT
        if bold_first and i == 0:
            run.bold = True
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])
    return row


def _make_section_title(doc, number_letter: str, title: str):
    """Agrega título de sección: 'I. TÍTULO'."""
    _add_formatted_paragraph(
        doc,
        f"{number_letter}. {title.upper()}:",
        bold=True, size=FONT_SIZE_SECTION,
        color=COLOR_PRIMARY,
        space_before=Cm(0.6), space_after=Cm(0.3),
        alignment=WD_ALIGN_PARAGRAPH.LEFT
    )


def _make_subsection_title(doc, letter: str, title: str):
    """Agrega subtítulo: 'A. TÍTULO'."""
    _add_formatted_paragraph(
        doc,
        f"{letter}. {title.upper()}",
        bold=True, size=FONT_SIZE_NORMAL,
        color=COLOR_SECONDARY,
        space_before=Cm(0.4), space_after=Cm(0.2),
        alignment=WD_ALIGN_PARAGRAPH.LEFT
    )


# ─── Secciones estáticas ──────────────────────────────────────────────────

SECCION_III_RESPONSABILIDADES = [
    "Ejecutar la coordinación, seguimiento y articulación de las actividades de los Equipos "
    "Básicos de Salud, garantizando la implementación del Plan de Cuidado Primario (PCP), el "
    "cumplimiento de metas del PICP y la integración con actores del sistema de salud en los territorios.",

    "Ejecutar actividades de planeación, organización, soporte operativo y gestión administrativa "
    "necesarias para la implementación de las acciones individuales, familiares y comunitarias "
    "de los Equipos Básicos de Salud.",

    "Ejecutar el proceso de facturación de los servicios generados por los Equipos Básicos de "
    "Salud, incluyendo la consolidación de soportes, validación de información, elaboración de "
    "RIPS, control de glosas y seguimiento a cuentas.",

    "Ejecutar la gestión, validación, depuración y reporte de la información generada en la "
    "operación de los EBS, asegurando consistencia, calidad del dato y cumplimiento de "
    "lineamientos técnicos del sector salud.",

    "Ejecutar actividades de administración de plataformas (SISPRO, PISIS, SIAPS), soporte "
    "técnico a usuarios, gestión de bases de datos, y mantenimiento de la infraestructura "
    "tecnológica asociada a la operación de los EBS.",

    "Ejecutar el análisis de datos, seguimiento de indicadores, generación de reportes técnicos "
    "y evaluación del cumplimiento de metas del PCP y PICP para apoyar la toma de decisiones.",

    "Ejecutar actividades de estructuración, revisión, seguimiento y control documental de los "
    "procesos contractuales, incluyendo elaboración de informes, cargue en plataformas "
    "(SECOP, SIA OBSERVA) y trazabilidad de la ejecución contractual.",
]

ACTIVIDADES_GENERALES = [
    ("1", "Apoyar procesos administrativos de la ESE",
     "1. Gestión de documentos, organización de información y apoyo operativo\n\n"
     "Se realizó la recepción, clasificación, organización y archivo de documentación relacionada "
     "con la operación de los Equipos Básicos en Salud (EBS), garantizando su integridad, "
     "disponibilidad y trazabilidad. Se brindó apoyo operativo en la gestión de flujos "
     "documentales, facilitando el acceso oportuno a la información por parte de las áreas requeridas."),

    ("2", "Apoyar seguimiento a actividades de los EBS",
     "2. Consolidación y revisión de información reportada por equipos\n\n"
     "Se efectuó la recopilación, depuración y consolidación de la información remitida por los "
     "equipos en territorio, verificando su coherencia, completitud y consistencia frente a los "
     "lineamientos institucionales. Se identificaron inconsistencias y se gestionaron ajustes con los responsables."),

    ("3", "Apoyar la elaboración de informes",
     "3. Construcción y consolidación de informes mensuales\n\n"
     "Se elaboraron informes administrativos mediante la integración de información operativa y "
     "documental, asegurando su estructuración técnica, claridad y correspondencia con los "
     "requerimientos de supervisión y seguimiento institucional."),

    ("4", "Gestionar información requerida por supervisión",
     "4. Recolección y organización de soportes para supervisión contractual\n\n"
     "Se gestionó la recopilación, verificación y organización de los soportes requeridos para "
     "la supervisión contractual, garantizando su adecuada disposición para procesos de revisión, "
     "validación y eventual auditoría."),

    ("5", "Apoyar la articulación entre áreas",
     "5. Coordinación entre equipos asistenciales y administrativos\n\n"
     "Se facilitó la articulación entre las áreas asistenciales y administrativas, mediante la "
     "gestión de comunicaciones, seguimiento a requerimientos y canalización de información, "
     "contribuyendo a la continuidad operativa de las actividades."),

    ("6", "Cumplir lineamientos institucionales",
     "6. Aplicación de directrices administrativas\n\n"
     "Se ejecutaron las actividades conforme a los lineamientos institucionales impartidos, "
     "asegurando su correcta implementación en los procesos administrativos y operativos."),

    ("7", "Apoyar procesos de facturación y cuentas de cobro (si aplica)",
     "7. Revisión preliminar y organización de documentos\n\n"
     "Se realizó la revisión inicial de documentos asociados a trámites administrativos "
     "(incluidas cuentas de cobro cuando aplicó), verificando requisitos básicos de forma y "
     "contenido antes de su remisión a las áreas competentes."),

    ("8", "Garantizar manejo adecuado de la información",
     "8. Custodia y confidencialidad de documentos\n\n"
     "Se garantizó el manejo adecuado de la información bajo criterios de reserva, "
     "confidencialidad y protección de datos, evitando accesos no autorizados o uso indebido "
     "de la información institucional."),

    ("9", "Entrega de informes y reportes",
     "9. Presentación oportuna de información solicitada\n\n"
     "Se atendieron los requerimientos de información de manera oportuna, asegurando la entrega "
     "dentro de los plazos establecidos y con calidad en el contenido suministrado."),

    ("10", "Apoyar actualización de bases de datos",
     "10. Actualización de información administrativa\n\n"
     "Se mantuvieron actualizadas las bases de datos y registros administrativos, garantizando "
     "consistencia, integridad y disponibilidad de la información para la toma de decisiones."),

    ("11", "Otras actividades asignadas",
     "11. Apoyo en tareas adicionales requeridas\n\n"
     "Se brindó apoyo en actividades complementarias asignadas por la supervisión o coordinación, "
     "en función de las necesidades del servicio, contribuyendo al cumplimiento integral de los "
     "objetivos institucionales."),
]


def _build_section_i(doc, mes: int, anio: int, periodo_texto: str = None):
    """Construye la Sección I: Objeto contractual y alcance."""
    mes_nombre = MESES_TITULO[mes]
    periodo = periodo_texto or f"{mes_nombre} de {anio}"

    _make_section_title(doc, "I", "OBJETO CONTRACTUAL Y ALCANCE")
    _add_formatted_paragraph(doc,
        f'En cumplimiento del Contrato No. 022 del 02 de febrero de 2026, cuyo objeto '
        f'corresponde a la prestación de servicios de apoyo administrativo para el '
        f'fortalecimiento de la Atención Primaria en Salud de la Empresa Social del Estado '
        f'Norte 3 E.S.E., en las unidades de atención de Puerto Tejada, Villa Rica y Padilla, '
        f'a través de la operación de los Equipos Básicos en Salud (EBS), conforme a la '
        f'Resolución 1010 de 2025, el presente informe describe la ejecución de actividades '
        f'desarrolladas durante el período comprendido entre el 01 y el 31 de {periodo}.'
    )
    _add_formatted_paragraph(doc,
        f'Para el período evaluado, la ejecución comprendió el acompañamiento a 16 Equipos '
        f'Básicos en Salud, distribuidos en las zonas urbanas y rurales priorizadas, conforme '
        f'a la planeación territorial establecida en el Plan de Cuidado Primario (PCP) y el '
        f'Plan Integral de Cuidado Primario (PICP).'
    )
    _add_formatted_paragraph(doc,
        'El alcance de las actividades desarrolladas se enmarcó en el apoyo a la gestión '
        'operativa, administrativa y de información, sin que ello implicara funciones de '
        'dirección o subordinación, garantizando la articulación funcional necesaria para la '
        'implementación de la estrategia de Atención Primaria en Salud.'
    )
    _add_formatted_paragraph(doc,
        'El alcance del contrato comprende el desarrollo de actividades de:',
        bold=True
    )
    for item in [
        "Coordinación operativa y seguimiento a la estrategia EBS",
        "Gestión administrativa y de información en salud",
        "Apoyo al proceso de facturación",
        "Administración de plataformas tecnológicas del sector salud",
        "Análisis de información y generación de reportes",
        "Apoyo a la gestión contractual",
    ]:
        _add_formatted_paragraph(doc, f"• {item}", size=FONT_SIZE_SMALL)
    _add_formatted_paragraph(doc,
        'Estas actividades se ejecutaron en articulación con los lineamientos definidos por '
        'la E.S.E. Norte 3, las entidades territoriales y el Ministerio de Salud y Protección '
        'Social, orientadas al cumplimiento de metas del Plan de Cuidado Primario (PCP) y del '
        'Plan Integral de Cuidado Primario (PICP).'
    )


def _build_section_ii(doc):
    """Construye la Sección II: Metodología de ejecución."""
    _make_section_title(doc, "II", "METODOLOGÍA DE EJECUCIÓN")
    _add_formatted_paragraph(doc,
        'La ejecución del contrato se desarrolló bajo un enfoque técnico-operativo '
        'estructurado en las siguientes fases:'
    )

    fases = [
        ("1. Planeación operativa", [
            "Elaboración y ajuste de cronogramas mensuales de actividades",
            "Definición de metas por territorio y microterritorio",
            "Organización de equipos de trabajo y asignación de responsabilidades",
        ]),
        ("2. Ejecución en territorio y soporte administrativo", [
            "Acompañamiento a los Equipos Básicos en Salud en la ejecución de actividades",
            "Desarrollo de acciones individuales, familiares y comunitarias",
            "Gestión operativa de jornadas extramurales",
        ]),
        ("3. Gestión de la información", [
            "Recolección, consolidación y depuración de bases de datos",
            "Validación de registros en plataformas (SISPRO, PISIS, SIAPS)",
            "Cruce de información entre fuentes asistenciales y administrativas",
        ]),
        ("4. Seguimiento y control", [
            "Monitoreo de indicadores de cobertura, oportunidad y calidad",
            "Verificación de cumplimiento de metas del PCP y PICP",
            "Control de calidad de la información reportada",
        ]),
        ("5. Facturación y soporte financiero", [
            "Consolidación de soportes para facturación",
            "Validación de registros RIPS",
            "Seguimiento a glosas, devoluciones y radicación de cuentas",
        ]),
        ("6. Generación de reportes", [
            "Elaboración de informes técnicos",
            "Consolidación de resultados operativos",
            "Entrega de información a supervisión, entidades territoriales y Ministerio",
        ]),
    ]

    for fase_title, items in fases:
        _add_formatted_paragraph(doc, fase_title, bold=True, space_before=Cm(0.2))
        for item in items:
            _add_formatted_paragraph(doc, f"• {item}", size=FONT_SIZE_SMALL)


def _build_section_iii(doc):
    """Construye la Sección III: Relación de perfiles."""
    _make_section_title(doc, "III", "RELACIÓN DE PERFILES DE LOS AFILIADOS PARTÍCIPES")
    _add_formatted_paragraph(doc,
        'El equipo de apoyo administrativo de los Equipos Básicos en Salud (EBS) de la E.S.E. '
        'Norte 3 está conformado por talento humano organizado funcionalmente para garantizar '
        'la ejecución de las actividades de coordinación operativa, apoyo administrativo, gestión '
        'de información, facturación, soporte tecnológico y gestión contractual, en el marco de '
        'la estrategia de Atención Primaria en Salud. En ese sentido, el equipo se estructura a '
        'partir de los siguientes perfiles funcionales:'
    )
    _add_formatted_paragraph(doc, "RESPONSABILIDAD PRINCIPAL", bold=True, space_before=Cm(0.3))
    for resp in SECCION_III_RESPONSABILIDADES:
        _add_formatted_paragraph(doc, resp, size=FONT_SIZE_SMALL)


def _build_section_iv(doc, apoyos: list[dict]):
    """Construye la Sección IV: Tabla de afiliados partícipes."""
    _make_section_title(doc, "IV", "INFORMACIÓN DE AFILIADOS PARTÍCIPES QUE EJECUTAN LAS ACTIVIDADES")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    col_widths = [1.0, 8.0, 3.5, 6.0]
    _add_table_header_row(table, ["NUM", "NOMBRE", "CÉDULA", "PERFIL"], col_widths)

    for i, ap in enumerate(apoyos, 1):
        bg = COLOR_TABLE_LIGHT_BG if i % 2 == 0 else None
        _add_data_row(table, [
            str(i),
            ap["nombre"],
            ap["identificacion"],
            ap["perfil"] or "",
        ], col_widths, bg_color=bg)

    _add_formatted_paragraph(doc, "", space_after=Cm(0.3))  # Espaciado después


def _build_section_v(doc):
    """Construye la Sección V: Actividades generales."""
    _make_section_title(doc, "V", "DESCRIPCIÓN DE CUMPLIMIENTO DE ACTIVIDADES GENERALES")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    col_widths = [1.0, 5.0, 12.0]
    _add_table_header_row(table, ["ÍTEM", "ACTIVIDAD ACORDADA", "ACTIVIDAD EJECUTADA"], col_widths)

    for item_num, actividad, ejecutada in ACTIVIDADES_GENERALES:
        _add_data_row(table, [item_num, actividad, ejecutada], col_widths)

    _add_formatted_paragraph(doc, "", space_after=Cm(0.3))


async def _build_section_vi(doc, db: AsyncSession, mes: int, anio: int,
                              apoyos: list[dict]):
    """Construye la Sección VI: Actividades específicas por perfil funcional."""
    _make_section_title(doc, "VI", "DESCRIPCIÓN DE CUMPLIMIENTO DE ACTIVIDADES ESPECÍFICAS")
    _add_formatted_paragraph(doc,
        'A continuación, se presentan las actividades específicas desarrolladas, organizadas '
        'por perfil funcional, en correspondencia con las obligaciones acordadas y las funciones '
        'asignadas en el marco de la ejecución del Contrato No. 022 del 02 de febrero de 2026.'
    )
    _add_formatted_paragraph(doc,
        'La siguiente descripción permite evidenciar la distribución de responsabilidades, la '
        'ejecución de las actividades y su articulación con el cumplimiento del objeto contractual, '
        'conforme a los lineamientos de la estrategia de Atención Primaria en Salud y la operación '
        'de los Equipos Básicos en Salud.',
        space_after=Cm(0.4)
    )

    # Letras para sub-secciones
    letras = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    for idx, ap in enumerate(apoyos):
        letra = letras[idx] if idx < len(letras) else f"Z{idx}"
        perfil_nombre = (ap["perfil"] or "APOYO ADMINISTRATIVO").upper()
        title = f"ACTIVIDADES DE {perfil_nombre}"
        if ap.get("nombres_extra"):
            title += f": {ap['nombre']} y {ap['nombres_extra']}"
        else:
            title += f": {ap['nombre']}"

        _make_subsection_title(doc, letra, title)

        # Obtener actividades de este apoyo
        result = await db.execute(
            select(ActividadApoyo)
            .where(ActividadApoyo.apoyo_id == ap["id"])
            .order_by(ActividadApoyo.orden)
        )
        actividades = result.scalars().all()

        if not actividades:
            _add_formatted_paragraph(doc,
                "No se registraron actividades específicas para este perfil en el período.",
                italic=True, size=FONT_SIZE_SMALL
            )
            continue

        # Crear tabla de actividades específicas
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        col_widths = [1.0, 6.0, 11.0]
        _add_table_header_row(table, ["ÍTEM", "ACTIVIDAD ACORDADA", "ACTIVIDAD REALIZADA"], col_widths)

        for act_idx, act in enumerate(actividades, 1):
            # Buscar evidencias para esta actividad
            ev_result = await db.execute(
                select(EvidenciaApoyo)
                .where(EvidenciaApoyo.actividad_apoyo_id == act.id)
                .order_by(EvidenciaApoyo.created_at.desc())
            )
            evidencias = ev_result.scalars().all()

            # Generar texto de actividad realizada basado en evidencias
            if evidencias:
                aprobadas = [e for e in evidencias if e.estado == "APROBADO"]
                pendientes = [e for e in evidencias if e.estado == "PENDIENTE"]
                rechazadas = [e for e in evidencias if e.estado == "RECHAZADO"]

                texto_realizada = f"{act_idx}. "

                if aprobadas:
                    texto_realizada += (
                        "Se ejecutó la actividad conforme a lo programado, presentando las "
                        "evidencias requeridas para su verificación y registro documental. "
                    )
                    texto_realizada += f"Se cuenta con {len(aprobadas)} evidencia(s) aprobada(s)."
                    # Detallar cada evidencia aprobada
                    for ev in aprobadas:
                        if ev.tipo == "TEXTO" and ev.contenido_texto:
                            texto_realizada += f"\n\n- Evidencia de texto: {ev.contenido_texto[:300]}"
                        elif ev.tipo == "IMAGEN":
                            nombre = ev.archivo_nombre or "imagen"
                            texto_realizada += f"\n\n- Evidencia gráfica: {nombre}"
                        elif ev.tipo == "ARCHIVO":
                            nombre = ev.archivo_nombre or "archivo"
                            texto_realizada += f"\n\n- Evidencia documental: {nombre}"
                elif rechazadas and not aprobadas:
                    texto_realizada += (
                        "La actividad se ejecutó parcialmente. Se presentaron soportes que "
                        "requieren ajustes de acuerdo con la revisión del coordinador. "
                    )
                else:
                    texto_realizada += (
                        "Se desarrollaron las acciones correspondientes a la actividad durante "
                        "el período, con soportes en proceso de revisión. "
                    )

                if rechazadas:
                    texto_realizada += f"\n{len(rechazadas)} evidencia(s) rechazada(s) requieren corrección."
            else:
                texto_realizada = (
                    f"{act_idx}. Actividad corresponde a las obligaciones contractuales del "
                    f"perfil. Pendiente de reporte de ejecución detallada."
                )

            _add_data_row(table, [
                str(act_idx),
                act.descripcion,
                texto_realizada,
            ], col_widths)

        _add_formatted_paragraph(doc, "", space_after=Cm(0.4))


# ─── Generador principal ──────────────────────────────────────────────────

async def generar_informe_apoyo(
    db: AsyncSession,
    mes: int,
    anio: int,
    periodo_texto: str = None,
) -> io.BytesIO:
    """Genera el informe mensual de apoyo administrativo en DOCX.

    Args:
        db: Sesión de base de datos.
        mes: Número del mes (1-12).
        anio: Año.
        periodo_texto: Texto opcional del período (ej: 'MAYO de 2026').

    Returns:
        BytesIO con el contenido del DOCX.
    """
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Obtener apoyos
    result = await db.execute(
        select(ApoyoAdministrativo)
        .where(ApoyoAdministrativo.activo == True)
        .order_by(ApoyoAdministrativo.nombre)
    )
    apoyos_db = result.scalars().all()
    apoyos = [{"id": a.id, "nombre": a.nombre, "identificacion": a.identificacion,
               "perfil": a.perfil} for a in apoyos_db]

    # ─── CARTA / ENCABEZADO ─────────────────────────────────────────────
    mes_nombre = MESES[mes]
    periodo = periodo_texto or f"01 al 31 de {MESES_TITULO[mes].lower()} de {anio}"

    # Logo (placeholder - se puede agregar después)
    _add_formatted_paragraph(doc, "Señor:", bold=False, space_after=Cm(0.1))
    _add_formatted_paragraph(doc, "DIDIER SANDOVAL BOTINA", bold=True, size=FONT_SIZE_NORMAL)
    _add_formatted_paragraph(doc, "Representante Legal", space_after=Cm(0.1))
    _add_formatted_paragraph(doc,
        'SINDICATO DE TRABAJADORES DE LA SALUD MÉDICA "SINTRASALUD MEDICA"',
        space_after=Cm(0.1))
    _add_formatted_paragraph(doc,
        "saludpublicapic2021@gmail.com",
        space_after=Cm(0.5))

    _add_formatted_paragraph(doc, "Cordial saludo,", space_after=Cm(0.3))

    _add_formatted_paragraph(doc,
        f'En cumplimiento del Contrato No. 022 del 02 de febrero de 2026, cuyo objeto '
        f'corresponde a la PRESTACIÓN DE SERVICIOS DE APOYO ADMINISTRATIVO PARA EL '
        f'FORTALECIMIENTO DE LA ATENCIÓN PRIMARIA EN SALUD DE LA EMPRESA SOCIAL DEL ESTADO '
        f'NORTE 3 ESE EN LAS UNIDADES DE ATENCIÓN DE PUERTO TEJADA VILLA RICA Y PADILLA '
        f'A TRAVÉS DE LA OPERACIÓN DE EQUIPOS BÁSICOS EN SALUD DE ACUERDO CON LA RESOLUCIÓN '
        f'1010 DE 2025, el equipo de apoyo administrativo presenta el SEGUNDO INFORME DE '
        f'ACTIVIDADES, correspondiente al período comprendido entre {periodo}.'
    )
    _add_formatted_paragraph(doc,
        'El presente informe da cuenta de las actividades desarrolladas en ejecución de las '
        'obligaciones contractuales, conforme a los lineamientos establecidos por la Entidad '
        'y en concordancia con las disposiciones aplicables a la operación de los Equipos '
        'Básicos en Salud, así como las actividades acordadas con el Sindicato.'
    )
    _add_formatted_paragraph(doc, "Atentamente,", space_before=Cm(0.5), space_after=Cm(0.5))
    _add_formatted_paragraph(doc, "EQUIPO DE APOYO ADMINISTRATIVO", bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add_formatted_paragraph(doc, "EQUIPOS BÁSICOS EN SALUD", bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add_formatted_paragraph(doc,
        'EMPRESA SOCIAL DEL ESTADO NORTE 3 -E.S.E.',
        bold=True, space_after=Cm(0.8),
        alignment=WD_ALIGN_PARAGRAPH.LEFT
    )

    # ─── SECCIONES I-VI ──────────────────────────────────────────────────
    _build_section_i(doc, mes, anio, periodo_texto)
    _build_section_ii(doc)
    _build_section_iii(doc)
    _build_section_iv(doc, apoyos)
    _build_section_v(doc)
    await _build_section_vi(doc, db, mes, anio, apoyos)

    # Guardar
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
