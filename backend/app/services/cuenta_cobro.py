"""Servicio para generar la Cuenta de Cobro en DOCX a partir de la plantilla
del backend (app/templates/cuenta_cobro.docx), rellenando los datos
contractuales del contratista y del contrato."""

import io
import logging
from datetime import date
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

DIAS_ES = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO", "DOMINGO"]
MESES_ES = [
    "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
    "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE",
]


def _set_para_text(parrafo, texto: str):
    """Reemplaza el texto de un párrafo preservando el formato del primer run."""
    if not parrafo.runs:
        parrafo.add_run(texto)
        return
    parrafo.runs[0].text = texto
    for r in parrafo.runs[1:]:
        r.text = ""


def _fecha_texto_es(fecha: date | None = None) -> str:
    """Formatea una fecha en español: 'MARTES 04 DE AGOSTO DE 2026'."""
    f = fecha or date.today()
    return f"{DIAS_ES[f.weekday()]} {f.day:02d} DE {MESES_ES[f.month - 1]} DE {f.year}"


def _valor_moneda(valor: float) -> str:
    """Formatea el valor con separador de miles y decimal en punto: 6500000 → '6.500.000.00'."""
    return f"{valor:,.2f}".replace(",", ".")


def generar_cuenta_cobro_docx(
    *,
    contratista_nombre: str,
    contratista_cedula: str,
    expedida_en: str | None,
    banco: str | None,
    tipo_cuenta: str | None,
    numero_cuenta: str | None,
    numero_contrato: str,
    objeto: str,
    valor: float,
    periodo_nombre: str,
    numero: str = "01",
    fecha: date | None = None,
) -> bytes:
    """Genera el DOCX de la cuenta de cobro a partir de la plantilla."""
    from app.services.numero_letras import numero_a_letras

    valor_letras = numero_a_letras(valor)
    if not valor_letras.upper().endswith("M/CTE"):
        valor_letras = f"{valor_letras} M/CTE"

    template_path = TEMPLATES_DIR / "cuenta_cobro.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Plantilla de cuenta de cobro no encontrada: {template_path}")

    doc = Document(str(template_path))
    paras = doc.paragraphs

    def p(idx: int):
        return paras[idx] if idx < len(paras) else None

    # 0 — Título
    _set_para_text(p(0), f"CUENTA DE COBRO N° {numero} - {periodo_nombre}")

    # 7 — Cuerpo principal: nombre, cédula, expedida en, suma en letras + valor
    expedida = f" expedida en {expedida_en}" if expedida_en else ""
    _set_para_text(
        p(7),
        f"{contratista_nombre} identificado con cédula de ciudadanía No. "
        f"{contratista_cedula}{expedida}, la suma de {valor_letras} "
        f"(${_valor_moneda(valor)}) por concepto de:",
    )

    # 9 — Concepto (objeto del contrato) + número de contrato
    _set_para_text(p(9), f"{objeto} de acuerdo con el contrato No. {numero_contrato}")

    # 12 — Fecha
    _set_para_text(p(12), f"Puerto Tejada Cauca, {_fecha_texto_es(fecha)}")

    # 17 — Firma (nombre)
    _set_para_text(p(17), contratista_nombre)

    # 18 — Firma (cédula)
    _set_para_text(p(18), f"Cédula de ciudadanía No. {contratista_cedula}{expedida}")

    # 19-21 — Datos bancarios
    _set_para_text(p(19), f"Banco: {banco or ''}")
    _set_para_text(p(20), f"Tipo de cuenta: {tipo_cuenta or ''}")
    _set_para_text(p(21), f"No. De cuenta: {numero_cuenta or ''}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
